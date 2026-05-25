from __future__ import annotations

from pathlib import Path
from typing import Any


def extract_resume_text(file_path: str | None) -> tuple[str | None, dict[str, Any]]:
    path_text = str(file_path or "").strip()
    if not path_text:
        return None, {"status": "failed", "reason": "missing_file_path"}
    path = Path(path_text).expanduser()
    if not path.exists() or not path.is_file():
        return None, {"status": "failed", "reason": "file_not_found", "file_path": path_text}
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".txt", ".md"}:
        return _normalize_text(path.read_text(encoding="utf-8", errors="ignore")), {
            "status": "completed",
            "engine": "plain-text",
            "file_type": suffix.lstrip("."),
        }
    return None, {"status": "unsupported", "reason": "unsupported_file_type", "file_type": suffix.lstrip(".") or "unknown"}


def _extract_pdf(path: Path) -> tuple[str | None, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        return None, {"status": "unavailable", "reason": "missing_dependency", "dependency": "pypdf"}
    reader = PdfReader(str(path))
    text = _normalize_text("\n\n".join(page.extract_text() or "" for page in reader.pages))
    return text, {"status": "completed" if text else "empty", "engine": "pypdf", "file_type": "pdf", "pages": len(reader.pages)}


def _extract_docx(path: Path) -> tuple[str | None, dict[str, Any]]:
    try:
        from docx import Document
    except ModuleNotFoundError:
        return None, {"status": "unavailable", "reason": "missing_dependency", "dependency": "python-docx"}
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = _normalize_text("\n".join(parts))
    return text, {"status": "completed" if text else "empty", "engine": "python-docx", "file_type": "docx"}


def _normalize_text(text: str | None) -> str | None:
    normalized = "\n".join(line.rstrip() for line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")).strip()
    return normalized or None
